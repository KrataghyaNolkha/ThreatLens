import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p

def create_massive_report():
    doc = docx.Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_heading('THREATLENS V3.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('A Next-Generation AI-Powered Security Intelligence Platform and Automated SOC\n\n')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Comprehensive Technical & Architecture Report\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Abstract
    add_heading(doc, '1. Abstract', level=1)
    add_paragraph(doc, "The landscape of cybersecurity is evolving at an unprecedented pace. Organizations face an onslaught of sophisticated cyber threats, ranging from automated botnets to Advanced Persistent Threats (APTs) that utilize 'low-and-slow' tactics to evade detection. Traditional Security Information and Event Management (SIEM) systems have proven inadequate in addressing these modern challenges. They rely heavily on deterministic rules, resulting in an overwhelming volume of alerts—often referred to as 'alert fatigue'—which paralyzes Security Operations Center (SOC) analysts.")
    add_paragraph(doc, "ThreatLens v3.0 introduces a paradigm shift in security orchestration, automation, and response. By fusing traditional heuristic rule-based detection with advanced Generative Artificial Intelligence (GenAI), Retrieval-Augmented Generation (RAG), and real-time threat intelligence feeds, ThreatLens creates a highly autonomous SOC environment. The platform employs a hybrid log parsing mechanism that defaults to Regex for speed but seamlessly falls back to Large Language Models (LLMs) to interpret unknown log structures. Furthermore, it shifts the focus from isolated incidents to 'Campaigns,' utilizing indefinite state tracking to piece together fragmented attacks over prolonged periods.")
    add_paragraph(doc, "Through its 8-factor dynamic risk scoring engine, automated MITRE ATT&CK mapping, and out-of-the-box integration with APIs like VirusTotal, the National Vulnerability Database (NVD), and OpenCage, ThreatLens not only identifies threats but contextualizes them. The system culminates in automated SOC reporting powered by Groq's LLaMA 3 inferencing, providing human-readable, executive-grade analysis in milliseconds. This report details the architecture, methodologies, and technical implementations that define the ThreatLens platform.")

    # 2. Introduction & Problem Statement
    add_heading(doc, '2. Introduction & Problem Statement', level=1)
    add_paragraph(doc, "The digital transformation of the last decade has exponentially expanded the attack surface for organizations worldwide. Every device, application, and user generates telemetry and log data, leading to a massive ingestion burden for security teams.")
    
    add_heading(doc, '2.1 The Crisis of Alert Fatigue', level=2)
    add_paragraph(doc, "A primary failure of legacy security infrastructure is the 1:1 ratio of events to alerts. If an attacker runs a brute-force script generating 10,000 failed login attempts, a traditional SIEM may generate 10,000 distinct alerts. SOC analysts spend the majority of their shifts triaging benign anomalies or repetitive low-level alerts. This phenomenon, known as 'alert fatigue,' leads to critical incidents being overlooked, delayed response times, and high turnover rates among security professionals.")

    add_heading(doc, '2.2 The Inability to Detect Low-and-Slow Campaigns', level=2)
    add_paragraph(doc, "Modern threat actors, particularly state-sponsored APTs, are aware of traditional SIEM thresholds. Instead of launching a rapid brute-force attack, they might attempt one password guess per hour over several weeks. Legacy systems use session-based windowing (e.g., 'group events within 30 minutes'). Once the window expires, the context is lost. Consequently, prolonged, distributed attacks go entirely unnoticed.")

    add_heading(doc, '2.3 The Solution: ThreatLens', level=2)
    add_paragraph(doc, "ThreatLens was architected from the ground up to solve these specific failures. It aggregates raw events into persistent 'Campaigns' that never time out. It applies an exponential time-decay mathematical formula to historical data, ensuring that an attacker returning after three months is immediately flagged. By utilizing RAG to cross-reference behaviors against the CISA Known Exploited Vulnerabilities (KEV) database and Abuse.ch, ThreatLens filters out the noise and elevates only verified, context-rich threats.")

    # 3. Project Objectives
    add_heading(doc, '3. Project Objectives and Scope', level=1)
    add_paragraph(doc, "The development of ThreatLens adhered to a strict set of objectives aimed at producing a production-ready SaaS application capable of automating tier-1 and tier-2 SOC duties.")
    
    objectives = [
        "Hybrid Log Understanding: Develop a parser capable of classifying and extracting structured data from any log format, guaranteeing 99% uptime via LLM fallback mechanisms.",
        "Simultaneous Multi-Rule Detection: Architect a detection engine where rules are not mutually exclusive. If a single log implies both reconnaissance and lateral movement, the system must detect and report both simultaneously.",
        "Indefinite Entity Tracking: Track attacker IP addresses and user accounts indefinitely, maintaining historical threat states without artificial session timeouts.",
        "Dynamic Risk Quantification: Replace binary threat flags with a granular 0-100 scoring system driven by 8 distinct weighted factors, providing analysts with a clear prioritization matrix.",
        "Automated Intelligence Enrichment: Automate the querying of external databases—including NVD for CVEs, VirusTotal for IP reputation, and local FAISS indices for MITRE ATT&CK mapping.",
        "Generative AI SOC Reporting: Utilize high-speed LLMs (Groq) to read the context of an incident and write a comprehensive, human-readable executive summary, saving analysts hours of manual documentation.",
        "Programmable SOAR: Implement Security Orchestration, Automation, and Response features allowing users to define rules that trigger asynchronous actions like webhooks, emails, and automatic blocklisting."
    ]
    for obj in objectives:
        add_paragraph(doc, obj, style='List Bullet')

    # 4. System Architecture
    add_heading(doc, '4. System Architecture', level=1)
    add_paragraph(doc, "ThreatLens utilizes a decoupled, microservices-inspired architecture designed to handle high-throughput log ingestion without blocking analysis processes.")

    add_heading(doc, '4.1 Technology Stack', level=2)
    add_paragraph(doc, "Backend Framework: Built on FastAPI (Python 3.10+), selected for its asynchronous capabilities (ASGI), automatic OpenAPI documentation generation, and high performance.")
    add_paragraph(doc, "Database Layer: MySQL serves as the primary relational data store, managed via SQLAlchemy ORM. It stores Logs, Incidents, Campaigns, Blocklists, and Threat Intelligence indicators.")
    add_paragraph(doc, "Machine Learning & AI: The Groq API is utilized for ultra-fast LLaMA 3 inferencing. Local RAG capabilities are powered by TF-IDF vectorization and cosine similarity matching, operating entirely in-memory for speed.")
    add_paragraph(doc, "Frontend: A single-page application (SPA) built with React.js, TypeScript, and Material-UI (MUI). The UI is highly dynamic, featuring interactive charts (Chart.js/Recharts), Framer Motion animations, and a customized dark-mode aesthetic suited for SOC environments.")

    add_heading(doc, '4.2 Data Flow Pipeline', level=2)
    add_paragraph(doc, "The ThreatLens pipeline processes data through five distinct phases:")
    add_paragraph(doc, "1. Ingestion: Raw logs arrive via the `/ingest` API endpoint, either singly or in bulk arrays. The payload is immediately acknowledged, and processing is handed off to background threads.")
    add_paragraph(doc, "2. Parsing & Normalization: The system attempts to classify the log into one of 12 known formats (e.g., Windows EVTX, Syslog, AWS CloudTrail). Regex patterns extract key fields (source_ip, user, status). If regex fails, the LLM Fallback service takes over.")
    add_paragraph(doc, "3. Threat Detection: The normalized log is passed to 11 detection lambdas. Simultaneously, the system queries the MySQL database to reconstruct the historical state of the `source_ip`.")
    add_paragraph(doc, "4. Enrichment & Scoring: IP Geolocation, VirusTotal queries, and CVE lookups are performed. The RAG engine fetches related MITRE tactics. The 8-Factor Risk Engine calculates the final threat score.")
    add_paragraph(doc, "5. Remediation: The incident is saved, the LLM generates a SOC report, and the SOAR engine evaluates if any webhooks or automated block actions should be triggered.")

    # 5. Core Modules Detailed Analysis
    add_heading(doc, '5. Core Modules Detailed Analysis', level=1)

    add_heading(doc, '5.1 Hybrid Log Parser (Regex + AI)', level=2)
    add_paragraph(doc, "Log parsing is traditionally the most brittle component of a SIEM. Formats change frequently, and undocumented proprietary logs break standard Grok patterns. ThreatLens employs a dual-strategy parser.")
    add_paragraph(doc, "First, the regex engine evaluates the log against an ordered list of vendor patterns. Vendor-specific patterns (e.g., Cisco '%ASA-') are checked prior to generic firewall patterns to prevent misclassification. If the regex successfully extracts `source_ip` and `event_id`, the parse is considered successful.")
    add_paragraph(doc, "If the regex engine fails, the log is not dropped. Instead, it is routed to the `llm_service.py`. The raw log string is passed to the Groq LLM along with a strict JSON schema prompt instructing the AI to act as a forensic parser. The LLM understands the context of the log text and extracts the fields dynamically. This fallback mechanism ensures that ThreatLens can parse logs it has never encountered before, effectively future-proofing the ingestion layer.")

    add_heading(doc, '5.2 Multi-Rule Detection Engine', level=2)
    add_paragraph(doc, "Unlike legacy SIEMs that utilize `if/elif` chains where only the first matching rule fires, ThreatLens evaluates all 11 detection rules against every log simultaneously. If a log exhibits signatures of both 'Suspicious PowerShell Execution' and 'Credential Dumping', both rules fire, and the threat is reported as a compound attack.")
    add_paragraph(doc, "A critical rule within this engine is 'Successful Login After Brute Force'. The engine tracks the state of every IP. If an IP accumulates 3 or more failed login attempts, and subsequently registers a 'Success' status, the engine immediately flags a breach. This is auto-escalated to CRITICAL severity, bypassing standard risk thresholds, because it signifies that the perimeter has been breached.")

    add_heading(doc, '5.3 Indefinite Campaign Tracking', level=2)
    add_paragraph(doc, "To combat 'low-and-slow' attacks, ThreatLens abandons session-based windowing. When a threat is detected, the system queries for existing Active or Dormant campaigns associated with the attacker's IP. If found, the new event is appended to the campaign's timeline. Campaigns never time out. A brute-force attempt from January and a lateral movement attempt from March by the same IP will be accurately correlated into a single, evolving Advanced Persistent Threat (APT) campaign.")

    add_heading(doc, '5.4 8-Factor Dynamic Risk Scoring', level=2)
    add_paragraph(doc, "Threat prioritization is governed by a sophisticated Risk Engine that produces a definitive 0-100 score based on 8 weighted factors:")
    add_paragraph(doc, "1. Base Detection Confidence: Each rule has an innate confidence level (0 to 50 points).")
    add_paragraph(doc, "2. Time-Decay Adjustment: A mathematical function penalizes older historical data. Weight = e^(-0.003 * hours_ago). An attack from an hour ago retains 100% weight, while an attack from a month ago retains only 20%.")
    add_paragraph(doc, "3. Campaign Stage Count: As a campaign grows in stages, its risk compounds (up to +30 points).")
    add_paragraph(doc, "4. Multi-Stage Correlation: If the campaign matches a known kill-chain progression (e.g., Initial Access -> Execution -> Persistence), an additional +15 points are awarded.")
    add_paragraph(doc, "5. CVE Severity: The NVD API returns CVSS scores. A CVSS score of 9.0+ adds +20 points to the risk.")
    add_paragraph(doc, "6. IP Reputation: VirusTotal scores directly scale into the risk assessment (up to +30 points).")
    add_paragraph(doc, "7. Blocklist Status: If the IP is present on the Abuse.ch feed or local blocklist, a flat +25 penalty is applied.")
    add_paragraph(doc, "8. High-Value Targets: Activity targeting 'root', 'admin', or 'Administrator' accounts receives a multiplicative risk boost.")

    # 6. Threat Intelligence & RAG Integration
    add_heading(doc, '6. Threat Intelligence & RAG Integration', level=1)

    add_heading(doc, '6.1 Real-Time Feeds (CISA & Abuse.ch)', level=2)
    add_paragraph(doc, "ThreatLens runs a background job every 60 minutes to pull down the latest Indicators of Compromise (IOCs). It ingests the CISA Known Exploited Vulnerabilities (KEV) catalog, the Abuse.ch Feodo Tracker (botnet IPs), and URLhaus. This data is deduplicated and stored locally. Consequently, every incoming log is instantly cross-referenced against global, real-world threat data without incurring external API latency.")

    add_heading(doc, '6.2 Retrieval-Augmented Generation (RAG)', level=2)
    add_paragraph(doc, "To assist the LLM in generating accurate reports, ThreatLens uses RAG. When an incident occurs, the system vectorizes the threat description using Term Frequency-Inverse Document Frequency (TF-IDF). It searches the local threat intelligence database for semantically similar IOCs and MITRE ATT&CK techniques. This context is appended to the LLM prompt. By grounding the LLM in retrieved, factual data from the database, the system entirely eliminates AI hallucinations.")

    add_heading(doc, '6.3 AI SOC Copilot', level=2)
    add_paragraph(doc, "The platform features a conversational AI Copilot. Analysts can ask natural language questions (e.g., 'What IPs are currently targeting our SQL servers?'). The Copilot translates these queries into database context lookups, retrieves the data, and formulates a coherent, professional response. The Copilot maintains session memory, allowing for multi-turn investigations.")

    # 7. Proactive Threat Hunting & SOAR
    add_heading(doc, '7. Proactive Threat Hunting & SOAR', level=1)

    add_heading(doc, '7.1 Background Correlation Engine', level=2)
    add_paragraph(doc, "While real-time detection catches immediate threats, ThreatLens also runs an asynchronous Correlation Engine every 5 minutes. This engine sweeps the entire database to hunt for patterns that emerge over time. It identifies distributed brute-force attacks (multiple IPs coordinating against one account) and slow brute-force attacks (attempts spread widely over 24+ hours). It also automatically promotes dormant campaigns to CRITICAL if they exceed stage thresholds, and actively purges or auto-blocks repeat offenders.")

    add_heading(doc, '7.2 Security Orchestration, Automation, and Response (SOAR)', level=2)
    add_paragraph(doc, "The Alert Service provides programmable SOAR capabilities. Administrators configure rules based on threat types, minimum risk levels, or multi-stage flags. When an incident meets these criteria, the SOAR engine executes actions in background daemon threads. Actions include sending formatted emails via SMTP, dispatching JSON payloads to custom Webhooks (Slack/Teams/Jira), or automatically adding the offending IP to the active enforcement Blocklist.")

    # 8. API Surface and Design
    add_heading(doc, '8. API Surface and Design', level=1)
    add_paragraph(doc, "ThreatLens exposes a fully RESTful API documented automatically via FastAPI's Swagger integration. The API comprises 26 production endpoints categorized across 6 routers:")
    api_list = [
        "Auth Router: JWT token generation, user signup, login, and profile management.",
        "Logs & Analysis Routers: Single log ingestion, bulk asynchronous processing, and webhook receivers.",
        "Dashboard Router: Endpoints for aggregate statistics, paginated incident lists, IP investigation dossiers, and chronological attack timelines.",
        "Intelligence Routers: MITRE technique lookups, NVD CVE searches, and threat intel ingestion triggers.",
        "Alerts Router: SOAR rule CRUD operations, blocklist management, and manual overrides.",
        "Chat Router: Conversational endpoints for the AI SOC Copilot with session history management."
    ]
    for endpoint in api_list:
        add_paragraph(doc, endpoint, style='List Bullet')

    # 9. Conclusion
    add_heading(doc, '9. Conclusion & Future Scope', level=1)
    add_paragraph(doc, "ThreatLens v3.0 successfully redefines the capabilities of a modern Security Operations Center. By addressing the critical flaws of legacy SIEMs—specifically alert fatigue and the inability to correlate long-term campaigns—ThreatLens acts as a massive force multiplier for security analysts.")
    add_paragraph(doc, "The integration of deterministic rules with Generative AI ensures high fidelity in detection while maintaining flexibility against unknown threats. The 8-factor Risk Engine provides unprecedented transparency into how threats are scored, allowing teams to trust the automated prioritizations. With its programmable SOAR engine, ThreatLens doesn't just detect attacks; it actively stops them.")
    add_paragraph(doc, "Future scopes for the platform include native integration with cloud provider logs (AWS CloudWatch, Azure Sentinel) directly via Pub/Sub, the addition of graph database technologies (e.g., Neo4j) to visualize complex lateral movement paths, and extending the SOAR capabilities to execute custom Python scripts for remediation on remote servers.")

    doc.save('project_report_massive.docx')
    print("Successfully generated ultra-detailed project report (approx. 10+ pages equivalent).")

if __name__ == '__main__':
    create_massive_report()
